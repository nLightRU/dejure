"""
    Docx generation module
"""

from typing import Tuple, Dict
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from database import Database

class Schedule:
    week_days_names = ('Понедельник', 'Вторник', 'Среда',
                           'Четверг', 'Пятница','Суббота',
                           'Воскресенье')
    def __init__(self, schedule_data, db: Database):
        self.schedule = schedule_data
        self.db = db
        self.days_num = 0
        self.has_orto = False
        
        for row in schedule_data:
            if row[0] in ('0', '1', '2', '3'):
                self.days_num += 1
            else:
                p = db.search_person(row)
                if not self.has_orto and 'орто' in p['job']:
                    self.has_orto = True
        
        months = ('', 
                  'январь', 'февраль', 'март',
                  'апрель', 'май', 'июнь',
                  'июль', 'август', 'сентябрь',
                  'октябрь', 'ноябрь', 'декабрь',
        )

        if schedule_data[0][3] == '0':
            month_num = int(schedule_data[0][4])
        else:
            month_num = int(schedule_data[0][3:4])
        self.month = months[month_num]
        self.year = int(schedule_data[0].split('.')[-1])

        self.pay_filename = f'{self.month}_{self.year}_график_дежурств_c_оплатой.docx'
        self.table_filename = f'{self.month}_{self.year}_график_дежурств.docx'
        

    def __create_table_row(self, person: Dict, date: str):
        name = person['last_name'] + ' ' + person['first_name'][0] + '.' + person['patronymic'][0] + '.'
        job = person['job'][0].upper() + person['job'][1:]
        job_short = self.db.get_short_job(person=person)
        return {
            'name': name, 
            'job':job,
            'job_short': job_short,
            'date': date
        }
    
    def __create_schedule_rows(self) -> tuple:
        """
            Selects persons from the database and 
            returns the tuple of dicts like 
            {
                'name':'Иванов А.И.', 
                'job':'Стоматолог-терапевт', 
                'job_short':'тер.', 
                'date':'22.01.2025'
            }
        """
        schedule_rows = []
        for line in self.schedule:
            if line[0] in ('0', '1', '2', '3'):
                date = line
                continue
            else:
                person_record = self.db.search_person(line)
                table_row = self.__create_table_row(person=person_record, date=date)
                schedule_rows.append(table_row)
        return tuple(schedule_rows)
    
    def __create_schedule_days(self) -> Tuple[Dict]:
        schedule_rows = self.__create_schedule_rows()
        dates = []
        days = []

        # creating list with duty dates without persons
        # list dates is needed for filtering same dates
        for schedule_row in schedule_rows:
            if schedule_row['date'] not in dates:
                dates.append(schedule_row['date'])
                days.append(
                    {
                        'date': schedule_row['date'],
                        'persons': []
                    }
                )
        
        # adding person to a specific day in days list
        for schedule_row in schedule_rows:
            for day in days:
                if schedule_row['date'] == day['date']:
                    s = f"{schedule_row['name']} - {schedule_row['job_short']}"
                    day['persons'].append(s)


        return tuple(days)

    def generate_schedule_with_pay(self, filedir='demo_pay.docx') -> None:
        schedule_rows = self.__create_schedule_rows()
            
        document = Document()
        
        normal_font = document.styles['Normal'].font
        normal_font.name = 'Times New Roman'
        normal_font.color.rgb = RGBColor(0, 0, 0)

        h_1_font = document.styles['Heading 1'].font
        h_1_font.size = Pt(14)
        h_1_font.bold = False
        h_1_font.name = 'Times New Roman'
        h_1_font.color.rgb = RGBColor(0, 0, 0)

        h_2_font = document.styles['Heading 2'].font
        h_2_font.size = Pt(12)
        h_2_font.name = 'Times New Roman'
        h_2_font.color.rgb = RGBColor(0, 0 , 0)

        table_signs = document.add_table(rows=1, cols=2)
        sign_cells = table_signs.rows[0].cells

        text_sign = 'СОГЛАСОВАНО: \n\n' + \
                'Председатель профсоюзной организации ТОГАУЗ ГСП №2 г. Тамбова\n\n' + \
                '____________ Е.П. Труфанова'

        
        p = sign_cells[0].add_paragraph(text_sign)
        

        text_sign = 'УТВЕРЖДАЮ:\n\n' + \
                    'Главный врач:\n\n' + \
                    '_______________ А.А. Корчагин'
        
        p = sign_cells[1].add_paragraph(text_sign)

        heading_text = 'График работы в выходные и/или нерабочие праздничные дни\n' + \
                        'ТОГАУЗ «Городская стоматологическая поликлиника №2 г. Тамбова»\n' + \
                        f'на {self.month} {self.year}\n' + \
                        'терапевтическое отделение №2 (Моршанское шоссе 22)'

        heading = document.add_heading(heading_text, level=1)
        heading.alignment = WD_TABLE_ALIGNMENT.CENTER

        table = document.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        
        hdr_cells[1].add_paragraph('ФИО')
        hdr_cells[2].add_paragraph('Должность')
        hdr_cells[3].add_paragraph('Дата')
        hdr_cells[4].add_paragraph('Оплата в двойном размере или выходной день')
        hdr_cells[5].add_paragraph('Согласие на работу в выходные и/или праздничные дни (подпись)')

        for cell in hdr_cells:
            cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            r = cell.paragraphs[0].add_run()
            r.bold = True

        # hdr_cells[0].style.font.bold = True 
        
        for i, sc_row in enumerate(schedule_rows, start=1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i) + '.'
            row_cells[1].text = sc_row['name']
            row_cells[2].text = sc_row['job']
            row_cells[3].text = sc_row['date']

            row_cells[2].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            row_cells[3].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER


        filepath = Path(filedir) / self.pay_filename

        p = document.add_paragraph("Зав. терапевтическим отделением №2 ____________________ Попов В.А.")

        if self.has_orto:
            p = document.add_paragraph("Зав. ортопедическим отделением ____________________ Момент О.А")

        document.save(filepath)

    def generate_schedule_table(self, filedir='demo_table.docx') -> None:
        """
            row_cells[0].add_paragraph('first', style='List Number')
        """

        document = Document()

        normal_font = document.styles['Normal'].font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(14)
        normal_font.color.rgb = RGBColor(0, 0, 0)

        h_1_font = document.styles['Heading 1'].font
        h_1_font.size = Pt(22)
        h_1_font.color.rgb = RGBColor(0, 0, 0)

        h_2_font = document.styles['Heading 2'].font
        h_2_font.size = Pt(16)
        h_2_font.name = 'Times New Roman'
        h_2_font.color.rgb = RGBColor(0, 0 , 0)


        heading_text = f'График дежурств на {self.month}'
        heading = document.add_heading(heading_text, level=1)
        heading.alignment = WD_TABLE_ALIGNMENT.CENTER
        

        if self.days_num % 2 == 0:
            t_rows = int(self.days_num / 2)
        else:
            t_rows = int(self.days_num / 2) + 1
        
        table = document.add_table(rows=t_rows, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        days = self.__create_schedule_days()
        text = f'test {self.days_num}'
        col_idx, row_idx = 0, 0
        for day in days:
            cell = table.cell(row_idx, col_idx)
            if col_idx == 1:
                row_idx += 1
                col_idx = 0
            else:
                col_idx = 1
            date_iso = '-'.join(reversed(day['date'].split('.')))
            week_day = date.fromisoformat(date_iso).weekday()
            week_day_name = Schedule.week_days_names[week_day]
            date_text = cell.add_paragraph(day['date'] + ' ' + week_day_name)
            date_text.style = document.styles['Heading 2']
            date_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
            start = 1
            for p in day['persons']:
                text = str(start) + '.' + '  ' + p
                cell.add_paragraph(text)
                start += 1
        
        
        p = document.add_paragraph("Зав. терапевтическим отделением №2 ____________________ Попов В.А.")

        if self.has_orto:
            p = document.add_paragraph("Зав. ортопедическим отделением ____________________ Момент О.А")

        for section in document.sections:
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)

        filepath = Path(filedir) / self.table_filename  
        document.save(filepath)
        

if __name__ == '__main__':
    print('docx generation module')
    
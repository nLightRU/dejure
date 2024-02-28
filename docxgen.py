"""
    Docx generation module
"""
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT

from dejure import schedule, search_person, create_table_row, create_schedule_rows

def generate_schedule_with_pay():
    schedule_rows = create_schedule_rows()
        
    document = Document()
    
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ФИО'
    hdr_cells[1].text = 'Должность'
    hdr_cells[2].text = 'Дата'

    # hdr_cells[0].style.font.bold = True 

    for i, sc_row in enumerate(schedule_rows, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = sc_row['name']
        row_cells[1].text = sc_row['job']
        row_cells[2].text = sc_row['date']

    document.save('demo_pay.docx')

def generate_schedule_table():

    document = Document()
    
    heading = document.add_heading('График дежурств', level=1)
    heading.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i in range(5):
        row_cells = table.rows[i].cells
        
        row_cells[0].add_paragraph('first', style='List Number')
        row_cells[0].add_paragraph('second', style='List Number')
        row_cells[0].add_paragraph('third', style='List Number')
        
        row_cells[1].add_paragraph('first', style='List Number')
        row_cells[1].add_paragraph('second', style='List Number')
        row_cells[1].add_paragraph('third', style='List Number')
        
        table.add_row()
        
    document.save('demo_table.docx')
        

if __name__ == '__main__':
    print('docx generation module')
    